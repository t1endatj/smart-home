from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Add Title
title = doc.add_heading('CHƯƠNG 5. KIẾN TRÚC PHẦN MỀM VÀ TÍCH HỢP TRỢ LÝ ẢO AI', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('5.1. Kiến trúc tổng thể hệ thống', level=2)
doc.add_heading('5.1.1. Kiến trúc 3 lớp của hệ thống', level=3)
doc.add_paragraph('Hệ thống Nhà thông minh được thiết kế theo kiến trúc 3 lớp tiêu chuẩn trong IoT, đảm bảo tính mô-đun hóa, dễ dàng mở rộng và bảo trì:')
doc.add_paragraph('Lớp Thiết bị (Perception Layer): Chịu trách nhiệm tương tác vật lý với môi trường. Bao gồm vi điều khiển ESP32 và các cảm biến (DHT11 đo nhiệt độ/độ ẩm, MQ-2 phát hiện khí gas, PIR phát hiện chuyển động), cùng các cơ cấu chấp hành (Servo cho cửa, Relay cho đèn và quạt).', style='List Bullet')
doc.add_paragraph('Lớp Xử lý Trung tâm (Edge/Server Layer): Đóng vai trò "não bộ" của hệ thống, được xây dựng trên nền tảng Python với FastAPI. Lớp này quản lý giao tiếp WebSocket thời gian thực, lưu trữ dữ liệu với SQLite, xử lý các logic tự động hóa (Automation) và tích hợp các dịch vụ AI.', style='List Bullet')
doc.add_paragraph('Lớp Giao diện (Application Layer): Được xây dựng bằng Vue.js 3. Cung cấp bảng điều khiển trực quan (Dashboard), mô hình Digital Twin 3D tương tác, biểu đồ giám sát theo thời gian thực và giao diện tương tác giọng nói, xác thực khuôn mặt.', style='List Bullet')

doc.add_heading('5.1.2. Chức năng của các thành phần hệ thống', level=3)
doc.add_paragraph('ESP32 Firmware (C++): Chạy vòng lặp chính để đọc dữ liệu cảm biến định kỳ, duy trì kết nối WebSocket tới Server. Xử lý phần cứng cấp thấp như tạo xung PWM để điều khiển góc quay Servo mượt mà (cơ chế hold pulse), bật/tắt Relay không gây xung đột.', style='List Bullet')
doc.add_paragraph('Backend Server (FastAPI): Cung cấp RESTful API cho Frontend tải dữ liệu lịch sử, đồng thời duy trì WebSocket Server để broadcast trạng thái (Delta State). Chứa Engine tự động hóa phân tích dữ liệu cảm biến để ra quyết định (ví dụ: phát hiện gas tự bật quạt 80% và mở cửa).', style='List Bullet')
doc.add_paragraph('Giao diện Frontend (Vue.js): Hiển thị trạng thái hệ thống, cho phép người dùng điều khiển thiết bị thông qua các nút bấm. Đồng bộ hóa hai chiều (Two-way sync) với Server để đảm bảo giao diện luôn phản ánh đúng trạng thái thực tế của phần cứng.', style='List Bullet')

doc.add_heading('5.1.3. Thiết kế cơ sở dữ liệu', level=3)
doc.add_paragraph('Hệ thống sử dụng SQLite làm cơ sở dữ liệu chính, được thiết kế nhỏ gọn với 2 bảng dữ liệu cốt lõi:')
doc.add_paragraph('Bảng sensor_logs: Lưu trữ dữ liệu lịch sử môi trường. Bao gồm các trường: id, temperature (nhiệt độ), humidity (độ ẩm), pir (chuyển động), gas_ppm (nồng độ gas), gas_alarm (cảnh báo gas), timestamp. Cho phép vẽ biểu đồ và phân tích xu hướng.', style='List Bullet')
doc.add_paragraph('Bảng home_state: Lưu trữ trạng thái toàn cục của hệ thống (Device States, Fan Speeds, Auto Configs). Chỉ duy trì duy nhất một bản ghi (id=1) được cập nhật liên tục dưới dạng chuỗi JSON (payload). Kèm theo trường revision để ngăn chặn xung đột dữ liệu (race condition) khi nhiều client cùng cập nhật.', style='List Bullet')

doc.add_heading('5.2. Thiết kế giao tiếp và truyền thông hệ thống', level=2)
doc.add_heading('5.2.1. Giao thức truyền thông giữa các thành phần', level=3)
doc.add_paragraph('Hệ thống kết hợp hai giao thức truyền thông chính: HTTP/REST cho các thao tác CRUD và tải dữ liệu một lần (One-off requests); và WebSocket cho truyền thông song công (Full-duplex) đảm bảo độ trễ thấp (Low-latency) cho lệnh điều khiển phần cứng.')

doc.add_heading('5.2.2. Thiết kế REST API', level=3)
doc.add_paragraph('Các API Endpoints quan trọng bao gồm:')
doc.add_paragraph('GET /api/state & POST /api/state: Đồng bộ và lưu trữ trạng thái toàn bộ ngôi nhà.', style='List Bullet')
doc.add_paragraph('POST /api/control: API điều khiển thiết bị đơn lẻ.', style='List Bullet')
doc.add_paragraph('GET /api/sensor/latest: Truy xuất dữ liệu cảm biến mới nhất cho Dashboard.', style='List Bullet')
doc.add_paragraph('POST /api/ai/command & POST /api/face/verify: Xử lý yêu cầu AI và nhận diện khuôn mặt.', style='List Bullet')

doc.add_heading('5.2.3. Cơ chế WebSocket thời gian thực', level=3)
doc.add_paragraph('Để tránh quá tải cho phần cứng ESP32 khi có quá nhiều lệnh, Server áp dụng cơ chế truyền tải trạng thái chênh lệch (Delta State). Khi có sự thay đổi trạng thái từ người dùng hoặc AI, Server sẽ so sánh Payload cũ và mới, sau đó chỉ gửi đi (broadcast) danh sách các lệnh thực sự bị thay đổi qua sự kiện device.sync. Điều này giúp ESP32 tối ưu hóa vòng lặp thực thi và không bị treo do xử lý lệnh dư thừa.')

doc.add_heading('5.3. Luồng dữ liệu và xử lý của hệ thống', level=2)
doc.add_heading('5.3.1. Luồng giám sát môi trường', level=3)
doc.add_paragraph('DHT11, MQ-2, PIR: ESP32 đọc tín hiệu từ các cảm biến, đóng gói thành chuỗi JSON và gửi qua WebSocket (sensor.sync) hoặc HTTP Post về Server định kỳ.')
doc.add_paragraph('Dashboard: Frontend gọi API GET /api/sensor/latest mỗi 5 giây để lấy dữ liệu môi trường mới nhất từ database và cập nhật giao diện trực quan.')

doc.add_heading('5.3.2. Luồng điều khiển thiết bị', level=3)
doc.add_paragraph('Dashboard → Server → ESP32: Khi người dùng bấm nút trên UI, Frontend gọi POST /api/control và lưu state lên Server. Server tính toán Delta State và phát qua WebSocket. ESP32 nhận lệnh, điều chỉnh GPIO (Relay) hoặc xung PWM (Servo).')
doc.add_paragraph('AI → Server → ESP32: Người dùng ra lệnh bằng giọng nói, Frontend gửi Text lên Server. Backend gọi API Gemini để phân tích, sau đó cập nhật State. Server tự động kích hoạt quá trình gửi WebSocket xuống ESP32.')

doc.add_heading('5.3.3. Luồng cảnh báo và phản hồi thời gian thực', level=3)
doc.add_paragraph('Gas Detection (Cảnh báo cháy nổ): Khi nồng độ Gas vượt ngưỡng (>2000ppm) hoặc MQ-2 kích hoạt mức thấp, ESP32 bật còi hú tại chỗ và báo lên Server. Server lập tức kích hoạt kịch bản SOS: Gửi lệnh đóng tất cả các cửa phòng, bật quạt hút bếp ở tốc độ cao (80%), bật toàn bộ hệ thống đèn thoát hiểm.')
doc.add_paragraph('Motion Detection: Cảm biến PIR phát hiện có người, Server (nếu bật chế độ Auto) sẽ tự động kích hoạt gửi lệnh bật đèn hành lang và quạt phòng khách.')

doc.add_heading('5.3.4. Sơ đồ trình tự xử lý (Sequence Diagram)', level=3)
doc.add_paragraph('1. (Khởi động) ESP32 kết nối WebSocket tới Server.')
doc.add_paragraph('2. ESP32 liên tục gửi bản tin Sensor Sync lên Server.')
doc.add_paragraph('3. Người dùng tương tác UI / Giọng nói gửi Request lên Server.')
doc.add_paragraph('4. Server lưu Database, tính toán Delta và phản hồi WebSocket.')
doc.add_paragraph('5. ESP32 nhận lệnh, xoay Servo hoặc bật Relay.')

doc.add_heading('5.4. Tích hợp trợ lý ảo AI', level=2)
doc.add_heading('5.4.1. Nhận diện khuôn mặt mở cửa tự động', level=3)
doc.add_paragraph('Hệ thống trang bị tính năng bảo mật bằng nhận diện khuôn mặt (Face ID) để mở khóa Dashboard. Trình duyệt sử dụng WebRTC getUserMedia luồng video, chụp ảnh Canvas và mã hóa base64 gửi lên /api/face/verify. Backend sử dụng thư viện xử lý ảnh để đối chiếu đặc trưng khuôn mặt (Face Encodings) với danh sách người dùng hợp lệ.')

doc.add_heading('5.4.2. Tương tác giọng nói hai chiều (STT/TTS)', level=3)
doc.add_paragraph('Frontend ứng dụng Web Speech API của trình duyệt để lắng nghe giọng nói theo thời gian thực (Speech-to-Text - STT). Dữ liệu văn bản được thu thập và chuyển giao trực tiếp lên Backend. Khi Backend xử lý xong và trả về phản hồi văn bản, giao diện sẽ in ra màn hình và có thể sử dụng Text-to-Speech (TTS) để máy giao tiếp lại với người dùng.')

doc.add_heading('5.4.3. Function Calling và Session Memory', level=3)
doc.add_paragraph('Lõi trí tuệ nhân tạo được vận hành thông qua API Google Gemini. Để AI có thể điều khiển được phần cứng, hệ thống cung cấp cho mô hình các cấu trúc hàm (Function Calling Schemas). AI sẽ không trả về văn bản tự do thuần túy, mà trả về một chuỗi JSON chuẩn hóa chứa mảng actions (danh sách thiết bị cần đổi trạng thái, tốc độ) và scenario (kịch bản ngữ cảnh như "sleep", "sos"). Cấu trúc này giúp Backend dễ dàng mapping lệnh ngôn ngữ tự nhiên (VD: "Trời nóng quá, bật quạt và mở cửa sổ") thành mã máy chính xác.')

doc.save('/home/phuchoangsrc/smart-home/Bao_cao_Chuong_5.docx')
